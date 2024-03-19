from docx import Document

def mostrar_partes(document):
    # Iterar sobre los párrafos del documento
    print("Párrafos:")
    for i, paragraph in enumerate(document.paragraphs):
        print(f"Párrafo {i + 1}: {paragraph.text}")

    # Iterar sobre las tablas del documento
    print("\nTablas:")
    for i, table in enumerate(document.tables):
        print(f"Tabla {i + 1}:")
        for row in table.rows:
            for cell in row.cells:
                print(cell.text, end="\t")
            print()

    # Iterar sobre las imágenes incrustadas en el documento
    print("\nImágenes:")
    for i, image in enumerate(document.inline_shapes):
        print(f"Imagen {i + 1}: {image.width}x{image.height}")

# Cargar el documento existente
doc = Document("documento.docx")

# Mostrar las partes del documento
mostrar_partes(doc)
